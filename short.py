import streamlit as st
import google.generativeai as genai
import os
from io import BytesIO
from pytube import YouTube
import tempfile
import requests
from urllib.parse import urlparse
import hashlib

# Configure Gemini API
GOOGLE_API_KEY = st.secrets["GOOGLE_API_KEY"]
genai.configure(api_key=GOOGLE_API_KEY)
model = genai.GenerativeModel('gemini-2.0-flash')

def extract_text_from_file(file):
    """Extracts text from various file types."""
    file_type = file.type
    if "text" in file_type:
        return file.getvalue().decode("utf-8")
    elif "pdf" in file_type:
        try:
            import pypdf
            pdf_reader = pypdf.PdfReader(BytesIO(file.getvalue()))
            text = ""
            for page_num in range(len(pdf_reader.pages)):
                text += pdf_reader.pages[page_num].extract_text() or ""
            return text
        except ImportError:
            st.error("Please install pypdf: `pip install pypdf`")
            return ""
    elif "docx" in file_type or "document" in file_type:
        try:
            import docx
            doc = docx.Document(BytesIO(file.getvalue()))
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        except ImportError:
            st.error("Please install python-docx: `pip install python-docx`")
            return ""
    else:
        return ""

# def download_youtube_transcript(youtube_url):
    # """Downloads and extracts transcript from a YouTube video."""
    # try:
        # yt = YouTube(youtube_url)
        # transcript = yt.captions.get_by_language_code('en')
        # if transcript:
            # return transcript.generate_srt_captions()
        # else:
            # return "No English transcript found."
    # except Exception as e:
        # return f"Error downloading YouTube transcript: {e}"

# def download_link_content(url):
    # """Downloads content from a given URL."""
    # try:
        # response = requests.get(url)
        # response.raise_for_status()
        # content_type = response.headers.get("content-type", "")
        # if "text" in content_type:
            # return response.text
        # elif "pdf" in content_type:
            # try:
                # import pypdf
                # pdf_reader = pypdf.PdfReader(BytesIO(response.content))
                # text = ""
                # for page_num in range(len(pdf_reader.pages)):
                    # text += pdf_reader.pages[page_num].extract_text() or ""
                # return text
            # except ImportError:
                # st.error("Please install pypdf: `pip install pypdf`")
                # return ""
        # elif "docx" in content_type or "document" in content_type:
            # try:
                # import docx
                # doc = docx.Document(BytesIO(response.content))
                # text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                # return text
            # except ImportError:
                # st.error("Please install python-docx: `pip install python-docx`")
                # return ""
        # else:
            # return "Content type not supported."
    # except requests.exceptions.RequestException as e:
        # return f"Error downloading content from URL: {e}"

def generate_content_hash(files):
    """Generates a hash based on file content, YouTube URL, and link URL."""
    content_hash = hashlib.sha256()
    if files:
        for file in files:
            content_hash.update(file.getvalue())
    # if youtube_url:
        # content_hash.update(youtube_url.encode('utf-8'))
    # if link_url:
        # content_hash.update(link_url.encode('utf-8'))

    return content_hash.hexdigest()

st.title("Gemini Content Q&A")

if 'all_content' not in st.session_state:
    st.session_state.all_content = ""
if 'content_hash' not in st.session_state:
    st.session_state.content_hash = ""

uploaded_files = st.file_uploader("Upload files (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"], accept_multiple_files=True)
# youtube_url = st.text_input("Enter YouTube URL")
# link_url = st.text_input("Enter a URL")

current_content_hash = generate_content_hash(uploaded_files)

if current_content_hash != st.session_state.content_hash:
    st.session_state.content_hash = current_content_hash
    st.session_state.all_content = ""
    if uploaded_files:
        for uploaded_file in uploaded_files:
            content = extract_text_from_file(uploaded_file)
            if content:
                st.session_state.all_content += content + "\n\n"

    # if youtube_url:
        # transcript = download_youtube_transcript(youtube_url)
        # if "Error" not in transcript:
            # st.session_state.all_content += transcript + "\n\n"
        # else:
            # st.error(transcript)

    # if link_url:
        # content = download_link_content(link_url)
        # if "Error" not in content and "Content type not supported" not in content:
            # st.session_state.all_content += content + "\n\n"
        # else:
            # st.error(content)

question = st.text_input("Ask a question about the uploaded content:")

if question and st.session_state.all_content:
    prompt = f"Answer the following question based only on the provided content.Never answer with information unrelated to the provided content:\n\nContent:\n{st.session_state.all_content}\n\nQuestion: {question}\n\nAnswer:"
    try:
        response = model.generate_content(prompt)
        st.write("Answer:")
        st.write(response.text)
    except Exception as e:
        st.error(f"Error generating response: {e}")
elif question and not st.session_state.all_content:
    st.warning("Please upload files, enter a YouTube URL, or a link before asking a question.")